# extractdocx.py a custom made document extractor for pdf and docx files

import os
import fitz  # PyMuPDF
import base64
import io
import re
import logging
from docx import Document
from PyPDF2 import PdfReader


# For cross-platform .doc to .docx conversion, you need 'unoconv' or a similar tool
# installed on your server. This is a more advanced setup. For now, we'll focus
# on what's possible with pure Python libraries where available.

# --- Configure Logging ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - [%(funcName)s] - %(message)s')
logger = logging.getLogger('doc_extractor')


def extract_content(file_path, contains_images=False):
    """
    Master function to extract content from a file based on its extension.
    This function routes the request to the appropriate specialized extractor.

    Args:
        file_path (str): The path to the input file.
        contains_images (bool): Flag to indicate if the PDF contains images.

    Returns:
        tuple: A tuple of (extracted_text, extracted_images_base64_list).
               'extracted_text' can be a string or a list of dicts.
               'extracted_images' is a list of base64 data URIs.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"The file was not found at path: {file_path}")

    file_extension = file_path.rsplit('.', 1)[-1].lower()
    logger.info(f"Extracting content from '{file_path}' (Type: {file_extension}, Images: {contains_images})")

    if file_extension == 'pdf':
        if contains_images:
            # If the user says it's an image-based PDF, we treat each page as an image.
            images = pdf_to_images(file_path)
            return None, images # Return only images
        else:
            # Otherwise, extract structured text and any embedded images.
            return extract_text_and_images_from_pdf(file_path)

    elif file_extension == 'docx':
        return extract_text_and_images_from_docx(file_path)
    
    elif file_extension == 'txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read(), [] # Return text and an empty list for images

    elif file_extension == 'doc':
        # Handling .doc is tricky without external dependencies like LibreOffice/unoconv.
        # This is a common limitation. We'll raise a clear error.
        raise NotImplementedError(
            "Direct extraction from .doc files is not supported due to its proprietary format. "
            "Please save the file as .docx and re-upload."
        )

    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


def extract_text_and_images_from_pdf(pdf_path):
    """
    Extracts both structured text and any embedded images from a PDF file.
    """
    content = []
    images = []
    try:
        with fitz.open(pdf_path) as doc:
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text("text")
                if text.strip():
                    content.append({"text": text, "formulas": extract_potential_formulas(text)})

                # Extract embedded image objects
                for img_index, img in enumerate(doc.get_page_images(page_num), 1):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    img_data = base_image["image"]
                    image_base64 = base64.b64encode(img_data).decode('utf-8')
                    images.append(f'data:image/{base_image["ext"]};base64,{image_base64}')
    except Exception as e:
        logger.error(f"Error processing PDF '{pdf_path}': {e}")
        return None, None
    return content, images


def extract_text_and_images_from_docx(docx_path):
    """
    Extracts text and embedded images from a .docx file directly.
    """
    doc = Document(docx_path)
    full_text = [para.text for para in doc.paragraphs]
    
    images = []
    # docx images are stored in 'inlines' or 'blips'
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_base64 = base64.b64encode(image_data).decode('utf-8')
            # You might need to determine the image type from headers if not obvious
            images.append(f'data:image/png;base64,{image_base64}') # Assuming PNG for simplicity

    return '\n'.join(full_text), images


def pdf_to_images(pdf_path):
    """
    Converts each page of a PDF into a high-quality image.
    Used when the entire page is treated as an image (e.g., scanned documents).
    """
    images = []
    try:
        with fitz.open(pdf_path) as doc:
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                # Use a higher DPI for better quality
                pix = page.get_pixmap(dpi=200)
                img_data = pix.tobytes("png")
                image_base64 = base64.b64encode(img_data).decode('utf-8')
                images.append(f'data:image/png;base64,{image_base64}')
    except Exception as e:
        logger.error(f"Error converting PDF pages to images: {e}")
        return []
    return images


def extract_potential_formulas(text):
    """
    Extracts potential LaTeX-style math formulas from text.
    """
    formula_pattern = r'\$[^$]+\$|\\\([^)]+\\\)|\\\[[^\]]+\\\]'
    return re.findall(formula_pattern, text)


def extract_text_from_any_file(file_path):
    """
    Utility to get plain text from PDF, DOCX, or TXT for answer key processing.
    """
    file_extension = file_path.rsplit('.', 1)[-1].lower()

    if file_extension == 'pdf':
        text_parts = []
        with open(file_path, 'rb') as f:
            pdf_reader = PdfReader(f)
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
        return "\n".join(text_parts)
    
    elif file_extension == 'docx':
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])

    elif file_extension == 'txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        raise ValueError(f"Cannot extract plain text from unsupported file type: {file_extension}")


def format_extracted_answers(answer_content_string):
    """
    Formats a string of answers (e.g., "1.A 2.B") into a standardized dictionary.
    """
    if not isinstance(answer_content_string, str):
        raise TypeError("Input to format_extracted_answers must be a string.")

    formatted_answers = {}
    # This regex is robust: finds a number, optional punctuation, and then a letter.
    pattern = re.compile(r'(\d+)\s*[.\s-]*\s*([A-Za-z])')
    
    for match in pattern.finditer(answer_content_string):
        question_number = f"q{match.group(1)}"  # Format as 'q1', 'q2', etc.
        answer = match.group(2).upper()
        formatted_answers[question_number] = answer

    return formatted_answers
